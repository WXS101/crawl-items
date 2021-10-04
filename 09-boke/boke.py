import requests
from lxml import etree
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
import time

title_list = []
data_list = []
href_list = []


def get_response(url):
    head = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36"
    }
    response = requests.get(url, headers=head).content.decode('utf-8')
    return response
    # if response.status_code == 200:
    #     return response.text
    # else:
    #     print("访问网页错误")


def parse_data(response):
    html = etree.HTML(response)
    datas = html.xpath('//*[@id="module_928"]/div[2]/div[1]/div[2]')
    for data in datas:
        for i in range(1, 51):
            title = data.xpath('./div[' + str(i) + ']/p[1]/span[2]/a/text()')
            href = data.xpath('./div[' + str(i) + ']/p[1]/span[2]/a/@href')[0]
            publish_time = data.xpath('./div[' + str(i) + ']/p[2]/span[2]/text()')
            article_title = " ".join(publish_time + title).replace(":", "点").replace("/", ".").replace("*", "").replace(
                "?", ".").replace("|", "").replace("\\", ",").replace("\"", ",").replace("<", ",").replace(">", ",")
            data_list.append([article_title, href])


def parse_details_data(response):
    data = etree.HTML(response)
    font = data.xpath('//*[@id="sina_keyword_ad_area2"]/font//text()')
    if font == []:
        font = ""
    else:
        font = font[0]
    text = data.xpath('//*[@id="sina_keyword_ad_area2"]//div//text()')
    # text2 = data.xpath('//*[@id="sina_keyword_ad_area2"]//p//text()')
    # text3 = data.xpath('//*[@id="sina_keyword_ad_area2"]//tr//text()')
    if text == [] or len(text) <= 5:
        text = data.xpath('//*[@id="sina_keyword_ad_area2"]//p//text()')
    if text == []:
        text = data.xpath('//*[@id="sina_keyword_ad_area2"]//tr//text()')
    else:
        pass
    # text = text1 + text2 + text3
    text = font + "".join(text).replace("&nbsp", "").replace("／", ",")
    return text


def main():
    # 随机延时
    # time.sleep(random.uniform(1, 4))
    # 代填入网址
    for i in range(84, 85):
        url = "http://blog.sina.com.cn/s/articlelist_1549639423_0_" + str(i) + ".html"
        response = get_response(url)
        parse_data(response)
    for i in range(50):
        print(i)
        response = get_response(data_list[i][1])
        time.sleep(0.5)
        text = parse_details_data(response)
        Doc = Document()
        Doc.styles['Normal'].font.name = u'宋体'
        Doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        Doc.add_paragraph(text)
        Doc.save(data_list[i][0] + ".docx")
        # file = docx.Document()
        # file.add_paragraph(text)
        # file.save(data_list[i][0] + ".docx")


if __name__ == '__main__':
    main()
